from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_plan_docx import (
    BLUE,
    BORDER,
    CONTENT_WIDTH_DXA,
    GOLD,
    INK,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    WHITE,
    add_body,
    add_bullet,
    add_callout,
    add_heading,
    add_hyperlink,
    add_table,
    apply_table_geometry,
    configure_styles,
    create_numbering,
    prevent_row_split,
    set_cell_borders,
    set_cell_shading,
    set_cell_text,
    set_repeat_table_header,
    set_run_font,
)


ROOT = Path(__file__).resolve().parent
BASE = ROOT / "salidas_proyecto_1_2026"
DATA = BASE / "datos"
RESULTS = BASE / "resultados"
DOCS = BASE / "documentacion"
REPORT = DOCS / "Informe_Final_Proyecto_1_MINEDUC.docx"
CODEBOOK = DOCS / "Libro_de_Codigos_MINEDUC.docx"
CODEBOOK_MD = DOCS / "Libro_de_Codigos_MINEDUC.md"

TEAM = [
    "ABBY SOFIA DONIS AGREDA — 22440",
    "PABLO DANIEL BARILLAS MORENO — 22193",
    "JORGE PALACIOS — 231385",
    "ROBERTO EMILIANO OTONIEL CAMPOSECO TORRES — 23968",
]


def fmt_int(value) -> str:
    return f"{int(value):,}"


def fmt_pct(value) -> str:
    return f"{float(value):.2f}%"


def add_field(paragraph, instruction: str, fallback="1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_header_footer(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run(short_title.upper())
    set_run_font(left, size=8.1, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("CC3084 · DATA SCIENCE · 2026")
    set_run_font(right, size=8.1, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("Página ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE", "1")
    r = fp.add_run(" de ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "NUMPAGES", "1")


def add_cover(doc: Document, title: str, subtitle: str, label: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSIDAD DEL VALLE DE GUATEMALA")
    set_run_font(r, size=11, color=GOLD, bold=True)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(title)
    set_run_font(r, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(subtitle)
    set_run_font(r, size=15, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run(label)
    set_run_font(r, size=10.5, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Integrantes")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    for member in TEAM:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(member)
        set_run_font(r, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Curso: CC3084 – Data Science · Semestre II 2026")
    set_run_font(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Versión final 2.0 · 17 de julio de 2026")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()


def add_figure(doc: Document, path: Path, caption: str, width=6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_run_font(r, size=8.8, color=MUTED, italic=True)


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    set_cell_borders(cell, color="CCD6DF", size=8, sides=("left",))
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for index, line in enumerate(code.strip().splitlines()):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.2, color="243447")
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 150, "end": 150},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_metric_cards(doc: Document, metrics: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=2, cols=len(metrics))
    fills = ["EAF2F8", "E8F5F2", "FFF4D6", "F1ECFA", "FDEEE8"]
    for i, (label, value) in enumerate(metrics):
        set_cell_shading(table.cell(0, i), fills[i % len(fills)])
        set_cell_shading(table.cell(1, i), fills[i % len(fills)])
        set_cell_text(table.cell(0, i), label, bold=True, color=MUTED, size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(1, i), value, bold=True, color=NAVY, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
        for r in range(2):
            set_cell_borders(table.cell(r, i), color=WHITE, size=10)
    widths = [CONTENT_WIDTH_DXA // len(metrics)] * len(metrics)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
        cell_margins_dxa={"top": 100, "bottom": 100, "start": 70, "end": 70},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_rule_card(doc: Document, variable: str, problem: str, rule: str, reason: str, risk: str, control: str) -> None:
    add_heading(doc, variable, 3)
    rows = [
        ("Problemas encontrados", problem),
        ("Regla ejecutada", rule),
        ("Justificación", reason),
        ("Riesgo y mitigación", risk),
        ("Control / evidencia", control),
    ]
    add_table(doc, ["Elemento", "Detalle"], rows, [2050, 7310], font_size=8.8, header_fill="EAF2F8")


def build_report() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    profile = pd.read_csv(DATA / "perfil_datos_crudos.csv")
    controls = pd.read_csv(DATA / "controles_calidad.csv")
    rules = pd.read_csv(DATA / "bitacora_reglas_resumen.csv")
    manifest = pd.read_csv(DATA / "manifiesto_archivos_crudos.csv")
    codebook = pd.read_csv(DATA / "libro_codigos.csv")
    sector = pd.read_csv(RESULTS / "distribucion_sector.csv")
    status = pd.read_csv(RESULTS / "distribucion_status.csv")
    area = pd.read_csv(RESULTS / "distribucion_area.csv")
    jornada = pd.read_csv(RESULTS / "distribucion_jornada.csv")
    modalidad = pd.read_csv(RESULTS / "distribucion_modalidad.csv")
    ubicacion = pd.read_csv(RESULTS / "distribucion_ubicacion.csv")
    telefono = pd.read_csv(RESULTS / "distribucion_telefono_estado.csv")
    calidad = pd.read_csv(DATA / "calidad_antes_despues.csv")
    transformaciones = pd.read_csv(DATA / "resumen_transformaciones.csv")
    plan_variables = pd.read_csv(DATA / "plan_limpieza_por_variable.csv")
    duplicados = pd.read_csv(DATA / "candidatos_duplicados_parciales.csv")
    cruzadas = pd.read_csv(DATA / "validaciones_consistencia_cruzada.csv")
    territorial = pd.read_csv(DATA / "catalogo_territorial_observado.csv")
    metadata = dict(pd.read_csv(DATA / "metadatos_conjunto.csv").astype(str).values)

    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc, "Proyecto 1 · Informe final")
    doc.core_properties.title = "Informe final: obtención y limpieza de datos MINEDUC"
    doc.core_properties.author = "; ".join(TEAM)
    doc.core_properties.subject = "Establecimientos educativos de Guatemala, nivel Diversificado"
    doc.core_properties.keywords = "MINEDUC, Data Science, limpieza, reproducibilidad, libro de códigos"

    add_cover(
        doc,
        "Informe final de obtención\ny limpieza de datos",
        "Establecimientos educativos de Guatemala · Nivel Diversificado",
        "PROYECTO 1 · ENTREGA FINAL",
    )

    add_heading(doc, "Resumen ejecutivo", 1)
    add_body(
        doc,
        "Se construyó un flujo reproducible para diagnosticar, limpiar, integrar y documentar los datos de establecimientos educativos autorizados por el Ministerio de Educación de Guatemala en el nivel Diversificado. El archivo crudo consolidado contiene 11,603 registros y 17 variables. Todas las columnas se cargaron como texto para impedir pérdidas de ceros iniciales en códigos, distritos y teléfonos.",
    )
    add_body(
        doc,
        "La ejecución conservó las 11,603 filas: CODIGO es único y no se detectaron duplicados exactos. La salida completa contiene 59 variables, incluidas columnas originales, versiones limpias, indicadores de ausencia, códigos nominales y metadatos. Se documentaron 13,031 pares candidatos a duplicado parcial sin eliminar registros automáticamente. Los 22 controles definidos finalizaron satisfactoriamente.",
    )
    add_metric_cards(doc, [("Registros", "11,603"), ("Crudas", "17"), ("Finales", "59"), ("Eliminadas", "0"), ("Controles", "22/22")])
    add_callout(
        doc,
        "Decisión central",
        "No se sobreescriben nombres ni etiquetas originales. La presentación corregida se guarda en columnas *_LIMPIO y la fuente permanece en *_ORIGINAL. Las claves sin tildes se usan solo para comparación, nunca para redactar informes.",
    )

    add_heading(doc, "Contenido del informe", 1)
    contents = [
        ("1", "Correspondencia con los lineamientos"),
        ("2", "Fuente, alcance y unidad de observación"),
        ("3", "Arquitectura reproducible y archivos crudos"),
        ("4", "Diagnóstico del estado inicial y duplicados parciales"),
        ("5", "Plan ejecutado y reglas por variable"),
        ("6", "Resultados y calidad antes/después"),
        ("7", "Análisis descriptivo del conjunto final"),
        ("8", "Control de calidad, consistencia y reproducibilidad"),
        ("9", "Libro de códigos, metadatos y entregables"),
        ("10", "Limitaciones, conclusiones y referencias"),
    ]
    add_table(doc, ["Sección", "Contenido"], contents, [1100, 8260], font_size=9.4)

    add_heading(doc, "1. Correspondencia con los lineamientos", 1)
    add_body(
        doc,
        "El trabajo se desarrolló con base en el documento “Proyecto 1. Obtención y Limpieza de datos” y en el buscador oficial del MINEDUC. La siguiente matriz muestra la evidencia concreta incluida en la entrega.",
    )
    compliance = [
        ("Carga correcta", "Notebook ejecutado; carga string; validación de 11,603 × 17 y esquema esperado.", "Cumplido"),
        ("Análisis crudo", "Perfil por variable: tipo, únicos, faltantes técnicos/semánticos, porcentaje, dominios, formatos y duplicados.", "Cumplido"),
        ("Limpieza y explicación", "Plan para 17 variables, bitácora por cambio, tabla de transformaciones, razones, riesgos y 22 pruebas.", "Cumplido"),
        ("Duplicados parciales", "TF–IDF de n-gramas, umbral 0.92, 13,031 pares, evidencia auxiliar y decisión por par; 0 eliminaciones automáticas.", "Cumplido"),
        ("Consistencia cruzada", "Siete verificaciones geográficas, de indicadores y de teléfonos, todas sin inconsistencias.", "Cumplido"),
        ("Calidad antes/después", "Doce métricas comparables, incluidos faltantes, formatos, duplicados y celdas transformadas.", "Cumplido"),
        ("Libro de códigos", "DOCX/PDF/Markdown y Excel con 59 variables, dominios, tratamientos, derivaciones, fecha, fuente y versión.", "Cumplido"),
        ("Conjunto limpio", "CSV nacional completo de 11,603 × 59 y vista analítica de 11,603 × 42.", "Cumplido"),
    ]
    add_table(doc, ["Criterio", "Evidencia", "Estado"], compliance, [1900, 6100, 1360], font_size=8.7)

    add_heading(doc, "2. Fuente, alcance y unidad de observación", 1)
    p = add_body(doc, "La fuente declarada es el buscador oficial de establecimientos educativos del Ministerio de Educación de Guatemala: ")
    add_hyperlink(p, "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/", "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/")
    add_body(
        doc,
        "El filtro de interés es Nivel Escolar = Diversificado. La unidad de observación es el código de servicio educativo autorizado; por ello, dos filas con nombres parecidos no necesariamente describen el mismo servicio. CODIGO se utiliza como llave y no el nombre del establecimiento.",
    )
    add_callout(
        doc,
        "Ciudad Capital",
        "El portal presenta Ciudad Capital como origen separado de Guatemala. Se conservaron 23 categorías de origen y se derivó una vista de 22 departamentos, en la que los 2,161 registros de Ciudad Capital se agregan a Guatemala. Las zonas capitalinas se mantienen en columnas separadas.",
        fill="FFF7E6",
        accent=GOLD,
    )
    add_heading(doc, "2.1 Metadatos del snapshot", 2)
    source_rows = [
        ("Fuente", metadata["FUENTE"]),
        ("URL", metadata["FUENTE_URL"]),
        ("Fecha de extracción", metadata["FECHA_EXTRACCION"]),
        ("Nota de la fecha", metadata["NOTA_FECHA_EXTRACCION"]),
        ("Fecha de procesamiento", metadata["FECHA_PROCESAMIENTO"]),
        ("Versión limpia", metadata["VERSION"]),
        ("SHA-256 del crudo", metadata["SHA256_CRUDO"]),
    ]
    add_table(doc, ["Metadato", "Valor"], source_rows, [2300, 7060], font_size=8.7)
    add_body(
        doc,
        "La fecha de extracción se toma del archivo crudo preservado. El CSV no incorporaba una hora interna de extracción, por lo que esta limitación se declara en lugar de inventar un timestamp. La huella SHA-256 permite comprobar que el snapshot no cambió durante el procesamiento.",
    )

    add_heading(doc, "2.2 Catálogo territorial y cobertura", 2)
    add_body(
        doc,
        "SEGEPLAN reporta 340 municipios para Guatemala en su conjunto territorial 2026. El snapshot MINEDUC contiene 334 códigos municipales con servicios de Diversificado y 22 zonas capitalinas. En consecuencia, el catálogo de este proyecto valida todas las combinaciones que aparecen en el universo analizado, pero no se presenta como listado nacional exhaustivo: seis municipios no aparecen en la consulta utilizada.",
    )
    territorial_summary = [
        ("Municipios observados", str(territorial.loc[territorial["TIPO_REGISTRO"].str.startswith("Municipio"), "MUNICIPIO_COD_FUENTE"].nunique())),
        ("Municipios nacionales de referencia", "340"),
        ("Zonas de Ciudad Capital", str((territorial["TIPO_REGISTRO"].str.startswith("Zona")).sum())),
        ("Departamentos analíticos", "22"),
        ("Origen del catálogo", "Snapshot MINEDUC; referencia SEGEPLAN 2026"),
    ]
    add_table(doc, ["Indicador", "Resultado"], territorial_summary, [4200, 5160], font_size=9)

    add_heading(doc, "3. Arquitectura reproducible y archivos crudos", 1)
    add_body(
        doc,
        "El notebook ejecutado concentra la carga, el diagnóstico, las funciones de normalización, la aplicación de reglas, la generación de salidas y las verificaciones. Cada ejecución vuelve a crear los archivos de datos, resultados y documentación a partir del CSV crudo.",
    )
    add_code(doc, """
df_raw = pd.read_csv(RAW_PATH, dtype="string", keep_default_na=False)
assert df_raw.shape == (11603, 17)
assert df_raw["CODIGO"].nunique() == 11603
assert int(df_raw.duplicated().sum()) == 0
""")
    add_body(
        doc,
        "El archivo consolidado se particionó por las 23 categorías geográficas publicadas para facilitar la auditoría. Estas particiones son derivaciones del consolidado crudo y no se presentan como descargas independientes. La suma de sus filas es 11,603.",
    )
    manifest_rows = [(r.ORIGEN, fmt_int(r.FILAS), r.ARCHIVO) for r in manifest.itertuples()]
    add_table(doc, ["Origen", "Filas", "Archivo crudo"], manifest_rows, [2700, 1000, 5660], font_size=8.2)

    add_heading(doc, "4. Diagnóstico del estado inicial", 1)
    add_heading(doc, "4.1 Dimensiones, tipos y unicidad", 2)
    add_body(
        doc,
        "El conjunto inicial tiene 11,603 filas y 17 variables. Las 17 se cargaron como string. Esta elección es semántica: CODIGO, DISTRITO y TELEFONO pueden contener dígitos, pero son identificadores y no magnitudes numéricas. Se encontraron 11,603 códigos únicos y 0 registros duplicados exactos.",
    )
    profile_rows = []
    for r in profile.itertuples():
        profile_rows.append((r.VARIABLE, r.TIPO_CARGADO, fmt_int(r.VALORES_UNICOS), f"{fmt_int(r.FALTANTES_ANALITICOS)} ({fmt_pct(r.PORCENTAJE_FALTANTE)})"))
    add_table(doc, ["Variable", "Tipo", "Valores únicos", "Faltantes analíticos"], profile_rows, [2300, 1450, 2000, 3610], font_size=8.3)

    add_heading(doc, "4.2 Valores faltantes", 2)
    add_body(
        doc,
        "Se distinguieron faltantes técnicos y ausencias semánticas. Los primeros son celdas vacías o marcadores sin contenido; las segundas son rótulos como SIN JORNADA y SIN ESPECIFICAR. En las columnas analíticas estos rótulos se transforman a null, pero la causa se conserva mediante el valor original y un indicador *_INFORMADA.",
    )
    add_figure(doc, RESULTS / "faltantes_analiticos.png", "Figura 1. Faltantes analíticos por campo, generados mediante código.")
    add_body(
        doc,
        "DIRECTOR registra 2,106 ausencias (18.15%), JORNADA 1,072 (9.24%), TELEFONO 933 (8.04%), SUPERVISOR 525 (4.52%), DISTRITO 522 vacíos técnicos, DIRECCION 87 y ESTABLECIMIENTO 5. Al validar el patrón del distrito se identificaron además 69 formatos incompletos; por ello DISTRITO_LIMPIO tiene 591 null.",
    )

    add_heading(doc, "4.3 Dominios, formatos e inconsistencias", 2)
    diagnostic_issues = [
        ("Identificadores", "CODIGO es íntegro y único; DISTRITO tiene vacíos e incompletos; TELEFONO contiene uno, varios o formatos históricos."),
        ("Geografía", "23 orígenes porque Ciudad Capital está separada; el análisis departamental requiere una derivación de 22 categorías."),
        ("Texto", "Mayúsculas sostenidas, espacios, signos, abreviaturas y tildes inconsistentes en nombres y direcciones."),
        ("Ausencias", "SIN JORNADA y SIN ESPECIFICAR describen ausencia, no categorías analíticas reales."),
        ("Categorías", "SECTOR, STATUS, MODALIDAD, AREA, JORNADA, PLAN y DEPARTAMENTAL requieren etiquetas legibles y códigos nominales estables."),
        ("Duplicados", "0 duplicados exactos y 0 códigos repetidos; nombres parecidos no justifican eliminar servicios."),
    ]
    add_table(doc, ["Grupo", "Problema potencial identificado"], diagnostic_issues, [1900, 7460], font_size=9)

    add_heading(doc, "4.4 Duplicados exactos y parciales", 2)
    add_body(
        doc,
        "No se encontraron duplicados exactos ni códigos repetidos. Para localizar coincidencias parciales se bloquearon los registros por departamento y municipio/zona y se comparó el nombre normalizado con TF–IDF de caracteres de 3 a 5 posiciones y distancia coseno. El umbral de similitud fue 0.92. Este método recupera errores de letras y variantes cercanas sin comparar indiscriminadamente registros de territorios distintos.",
    )
    duplicate_counts = duplicados["DECISION"].value_counts()
    duplicate_rows = [
        ("Pares candidatos", fmt_int(len(duplicados)), "Todos tienen decisión documentada"),
        ("Conservar: servicios distintos", fmt_int(duplicate_counts.get("CONSERVAR_SERVICIOS_DISTINTOS", 0)), "Cambian sector, área, modalidad, jornada, plan o departamental"),
        ("Conservar y revisar", fmt_int(duplicate_counts.get("CONSERVAR_Y_REVISAR", 0)), "Coinciden servicio y dirección o teléfono, pero los códigos son distintos"),
        ("Conservar códigos distintos", fmt_int(duplicate_counts.get("CONSERVAR_CODIGOS_DISTINTOS", 0)), "No existe evidencia suficiente para fusionar"),
        ("Revisión manual recomendada", fmt_int(duplicados["REVISION_MANUAL_RECOMENDADA"].sum()), "Requiere contraste con la fuente oficial antes de una eventual fusión"),
        ("Filas eliminadas", "0", "La similitud de nombre nunca autoriza eliminación automática"),
    ]
    add_table(doc, ["Resultado", "Cantidad", "Criterio"], duplicate_rows, [2700, 1500, 5160], font_size=8.5)
    add_callout(
        doc,
        "Regla conservadora",
        "CODIGO identifica un servicio autorizado. Dos nombres iguales pueden corresponder a jornadas, planes, sectores o códigos diferentes. Por esa razón, el archivo de candidatos documenta la revisión, pero el conjunto final conserva todas las filas.",
        fill="FFF7E6",
        accent=GOLD,
    )

    add_heading(doc, "5. Plan ejecutado y reglas por variable", 1)
    add_body(
        doc,
        "Las reglas se aplicaron de forma conservadora. Unicode NFC, recorte de espacios y presentación ortográfica se ejecutaron en columnas derivadas. Las correcciones léxicas se limitaron a casos respaldados; las claves sin tildes no sustituyen el texto de presentación. A continuación se documentan problema, regla, razón, riesgo y control para cada variable cruda.",
    )

    cards = {
        "CODIGO": (
            "No presenta faltantes ni repeticiones; el riesgo principal es que el software lo convierta a número o fecha.",
            "Cargar y exportar como texto; comprobar unicidad y patrón; conservarlo como llave del servicio.",
            "Los guiones y ceros forman parte del identificador.",
            "Conversión automática. Se evita fijando dtype='string' y validando la salida.",
            "11,603 valores válidos y únicos; QC02 y QC03 aprobados.",
        ),
        "DISTRITO": (
            "522 vacíos técnicos y 69 formatos incompletos; 1,667 valores distintos en la fuente.",
            "Conservar DISTRITO_ORIGINAL; producir DISTRITO_LIMPIO solo si el formato está completo y clasificar la celda como Válido, Vacío o Incompleto.",
            "No debe inventarse el fragmento faltante de un código administrativo.",
            "Perder un prefijo parcial. Se mantiene original y estado para una revisión posterior.",
            "591 null en DISTRITO_LIMPIO, exactamente los esperados; QC08 aprobado.",
        ),
        "DEPARTAMENTO": (
            "La fuente tiene 23 orígenes porque Ciudad Capital se publica fuera de Guatemala.",
            "Crear etiqueta legible, conservar DEPARTAMENTO_ORIGINAL, derivar DEPARTAMENTO_ANALISIS de 22 categorías y agregar código territorial de dos dígitos.",
            "Evita doble conteo y permite análisis nacional coherente sin perder procedencia.",
            "Fusionar orígenes. La columna original y UBICACION_GRUPO mantienen la distinción.",
            "23 orígenes, 22 departamentos analíticos y 2,161 filas capitalinas; QC05–QC07 aprobados.",
        ),
        "MUNICIPIO": (
            "Combina municipios y zonas capitalinas; nombres en mayúsculas y tildes no uniformes.",
            "Crear MUNICIPIO_LIMPIO, TIPO_UBICACION, ES_CIUDAD_CAPITAL, ZONA_CAPITAL_COD y MUNICIPIO_COD_FUENTE.",
            "Separa la semántica de zona y municipio y facilita agrupaciones posteriores.",
            "Tratar una zona como municipio. El tipo de ubicación se conserva explícitamente.",
            "Las 350 etiquetas crudas permanecen trazables; 22 zonas capitalinas identificadas.",
        ),
        "ESTABLECIMIENTO": (
            "6,170 valores únicos, 5 vacíos, mayúsculas sostenidas, tildes y algunos errores evidentes.",
            "Conservar el original; crear presentación española conservadora y una clave auxiliar sin tildes/puntuación solo para comparar candidatos. No deduplicar por similitud.",
            "Los informes necesitan nombres legibles, pero los servicios se distinguen por CODIGO.",
            "Corregir indebidamente un nombre propio. Solo se aplican reglas respaldadas y se guarda el original.",
            "5 null; ninguna fila eliminada; R05 y QC09 documentan el resultado.",
        ),
        "DIRECCION": (
            "7,260 valores únicos, 87 vacíos, abreviaturas, espacios y puntuación heterogéneos.",
            "Normalizar espacios y puntuación de presentación; preservar abreviaturas plausibles; asignar null solo a vacío real o signos aislados.",
            "Mejora la lectura sin inventar componentes geográficos.",
            "Alterar una referencia válida. DIRECCION_ORIGINAL permanece disponible.",
            "87 null en DIRECCION_LIMPIA; QC10 aprobado.",
        ),
        "TELEFONO": (
            "933 vacíos; celdas con uno o varios números y 104 formatos no estándar.",
            "Tratar como texto, normalizar candidatos de ocho dígitos, extraer TELEFONO_PRINCIPAL, clasificar TELEFONO_ESTADO y dejar null en TELEFONO_LIMPIO cuando no sea interpretable.",
            "Un teléfono no es una cantidad y no debe imputarse.",
            "Un formato histórico utilizable podría no satisfacer la regla. TELEFONO_ORIGINAL y el estado conservan la evidencia.",
            "10,444 válidos únicos, 122 múltiples, 104 no estándar, 933 vacíos y 1,037 null analíticos.",
        ),
        "SUPERVISOR": (
            "1,268 valores únicos y 525 ausencias; nombres en mayúsculas y tildes variables.",
            "Crear presentación conservadora; null solo para ausencia real; conservar original.",
            "Permite lectura institucional sin afirmar identidad entre homónimos.",
            "Sobrecorregir apellidos o fusionar personas. No se asignan identificadores personales.",
            "525 null en SUPERVISOR_LIMPIO; QC12 aprobado.",
        ),
        "DIRECTOR": (
            "5,397 valores únicos y 2,106 ausencias, incluidos marcadores como '--'.",
            "Convertir marcadores vacíos a null en DIRECTOR_LIMPIO, crear estado y aplicar presentación conservadora.",
            "Distingue ausencia de nombre informado.",
            "Inferir una persona o alterar apellidos. Se preserva DIRECTOR_ORIGINAL.",
            "2,106 null y estado explícito; QC13 aprobado.",
        ),
        "NIVEL": (
            "Variable constante con una categoría: DIVERSIFICADO.",
            "Normalizar la etiqueta a Diversificado y verificar que no aparezcan otros niveles.",
            "Confirma el alcance definido por el filtro de extracción.",
            "Ocultar una descarga mezclada. El control exige una sola categoría.",
            "11,603 filas en Diversificado.",
        ),
        "SECTOR": (
            "Cuatro categorías válidas, originalmente en mayúsculas.",
            "Crear SECTOR_LIMPIO y SECTOR_COD con catálogo nominal 1–4; conservar etiqueta original.",
            "Los códigos estabilizan análisis y filtros sin reemplazar las etiquetas.",
            "Interpretar los números como orden. El catálogo declara 'Código nominal'.",
            "0 códigos faltantes; QC16 aprobado.",
        ),
        "AREA": (
            "URBANA, RURAL y tres filas SIN ESPECIFICAR.",
            "Crear AREA_LIMPIA y AREA_COD; representar SIN ESPECIFICAR como null y AREA_INFORMADA=0; conservar original.",
            "SIN ESPECIFICAR no es una tercera área real.",
            "Perder la causa del null. El indicador y la columna original la preservan.",
            "3 null en AREA_LIMPIA; QC14 aprobado.",
        ),
        "STATUS": (
            "Cinco categorías administrativas legibles pero en mayúsculas.",
            "Crear STATUS_LIMPIO y STATUS_COD con catálogo nominal.",
            "Facilita comparación y mantiene etiquetas aptas para informes.",
            "Confundir código con jerarquía. Se documenta como nominal.",
            "0 códigos faltantes; QC17 aprobado.",
        ),
        "MODALIDAD": (
            "Dos categorías; requiere presentación ortográfica con diéresis y tilde.",
            "Crear MODALIDAD_LIMPIA y MODALIDAD_COD; conservar original.",
            "Evita variantes de escritura en resultados.",
            "Eliminar diacríticos. Las etiquetas finales son Monolingüe y Bilingüe.",
            "11,130 monolingües y 473 bilingües; catálogo completo.",
        ),
        "JORNADA": (
            "Seis valores crudos; 1,072 filas dicen SIN JORNADA.",
            "Crear JORNADA_LIMPIA y código; transformar SIN JORNADA a null, marcar JORNADA_INFORMADA=0 y conservar original.",
            "Una ausencia no debe competir con jornadas reales.",
            "Perder la razón del null. Original e indicador permanecen.",
            "1,072 null en JORNADA_LIMPIA; QC15 aprobado.",
        ),
        "PLAN": (
            "Trece categorías con paréntesis, acentos y capitalización heterogénea.",
            "Crear PLAN_LIMPIO y PLAN_COD mediante catálogo nominal; conservar original.",
            "Evita divergencias posteriores en filtros y tablas.",
            "Asignar un orden inexistente. El código es estrictamente nominal.",
            "13 categorías documentadas y sin códigos faltantes.",
        ),
        "DEPARTAMENTAL": (
            "26 denominaciones administrativas; texto en mayúsculas.",
            "Crear DEPARTAMENTAL_LIMPIO y código nominal con catálogo; conservar original.",
            "Permite agrupar oficinas sin perder su rótulo institucional.",
            "Confundir oficina administrativa con departamento geográfico. Ambas variables permanecen separadas.",
            "26 etiquetas documentadas y códigos completos.",
        ),
    }
    for variable in profile["VARIABLE"]:
        add_rule_card(doc, variable, *cards[variable])

    add_heading(doc, "5.1 Resumen de reglas y filas afectadas", 2)
    rule_rows = [(r.REGLA_ID, r.VARIABLES, r.ACCION, fmt_int(r.FILAS_AFECTADAS)) for r in rules.itertuples()]
    add_table(doc, ["ID", "Variables", "Acción", "Filas afectadas"], rule_rows, [700, 1800, 5460, 1400], font_size=7.9)

    add_heading(doc, "5.2 Tabla de transformaciones exigida", 2)
    add_body(
        doc,
        "La siguiente tabla usa exactamente las cinco columnas solicitadas por la guía 2026. Registros afectados cuenta las celdas cuyo valor analítico difiere del valor crudo, incluido el paso explícito a null cuando corresponde.",
    )
    transformation_rows = [
        (r["Variable"], r["Problema detectado"], r["Transformación"], fmt_int(r["Registros afectados"]), r["Justificación"])
        for r in transformaciones.to_dict("records")
    ]
    add_table(
        doc,
        ["Variable", "Problema detectado", "Transformación", "Registros afectados", "Justificación"],
        transformation_rows,
        [1050, 2150, 2500, 1200, 2460],
        font_size=6.9,
        header_fill="EAF2F8",
    )

    add_heading(doc, "6. Resultados de la limpieza", 1)
    add_body(
        doc,
        "La tasa de conservación fue 100%: (11,603 / 11,603) × 100. No se eliminaron registros, porque la llave CODIGO es íntegra y no existen duplicados exactos. Las transformaciones se expresan como columnas nuevas y una bitácora de cambios registra, por fila y variable, el valor original, el valor limpio y la regla aplicada.",
    )
    add_heading(doc, "6.1 Calidad antes y después", 2)
    quality_rows = []
    for r in calidad.to_dict("records"):
        def show(value):
            number = float(value)
            return f"{number:,.4f}" if not number.is_integer() else f"{int(number):,}"
        quality_rows.append((r["METRICA"], show(r["ANTES"]), show(r["DESPUES"]), show(r["CAMBIO"])))
    add_table(doc, ["Métrica", "Antes", "Después", "Cambio"], quality_rows, [4960, 1400, 1500, 1500], font_size=8.0)
    add_body(
        doc,
        "El total de null analíticos aumenta en 173 celdas (de 5,253 a 5,426). Esto no es un deterioro: 69 distritos incompletos y 104 teléfonos no interpretables dejan de aparentar ser valores válidos. Sus textos originales permanecen disponibles. A la vez, las tres variables con formatos inconsistentes quedan en cero y los 13,031 pares candidatos pasan de no tener decisión a contar con una decisión documentada.",
    )
    add_heading(doc, "6.2 Ejemplos de transformación", 2)
    examples = [
        ("DEPARTAMENTO", "ALTA VERAPAZ", "Alta Verapaz"),
        ("MUNICIPIO", "COBAN", "Cobán"),
        ("ESTABLECIMIENTO", "INSTITUTO MIXTO NOCTURNO FRANCISCO MARROQUIN", "Instituto Mixto Nocturno Francisco Marroquín"),
        ("ESTABLECIMIENTO", "INSTITUTO DE TURSMO Y AVIACON DEL NORTE I.T.A.N", "Instituto de Turismo y Aviación del Norte I.T.A.N"),
        ("DIRECCION", "6A. AVENIDA 1-15 ZONA 4", "6A. Avenida 1-15, Zona 4"),
        ("TELEFONO", "77945104", "7794-5104"),
        ("SUPERVISOR", "JORGE EDUARDO PAQUE LAZARO", "Jorge Eduardo Paque Lázaro"),
        ("TELEFONO", "formato no interpretable", "null; original y estado conservados"),
        ("DIRECTOR", "--", "null; estado = Marcador"),
    ]
    add_table(doc, ["Variable", "Fuente", "Presentación analítica"], examples, [1700, 3550, 4110], font_size=8.4)
    add_callout(
        doc,
        "Interpretación",
        "Estas correcciones no certifican que cada nombre propio coincida con una fuente externa; son normalizaciones reproducibles respaldadas por reglas conservadoras. Para usos legales o administrativos debe verificarse el valor con MINEDUC y consultar siempre la columna original.",
        fill="FFF7E6",
        accent=GOLD,
    )

    add_heading(doc, "7. Análisis descriptivo del conjunto final", 1)
    add_figure(doc, RESULTS / "resumen_distribuciones.png", "Figura 2. Distribuciones principales del conjunto limpio.")
    add_heading(doc, "7.1 Sector y estado", 2)
    sector_rows = [(r.CATEGORIA, fmt_int(r.FILAS), fmt_pct(r.PORCENTAJE)) for r in sector.itertuples()]
    status_rows = [(r.CATEGORIA, fmt_int(r.FILAS), fmt_pct(r.PORCENTAJE)) for r in status.itertuples()]
    add_table(doc, ["Sector", "Filas", "%"], sector_rows, [5200, 1900, 2260], font_size=9)
    add_table(doc, ["Estado", "Filas", "%"], status_rows, [5200, 1900, 2260], font_size=9)
    add_body(
        doc,
        "El 83.15% de los servicios pertenece al sector privado. El 57.76% figura abierto; el 25.25% está cerrado temporalmente y el 15.63% cerrado definitivamente. Estas cifras describen registros administrativos del archivo y no deben interpretarse como matrícula, capacidad instalada ni edificios únicos.",
    )

    add_heading(doc, "7.2 Área, jornada y modalidad", 2)
    arows = [(r.CATEGORIA, fmt_int(r.FILAS), fmt_pct(r.PORCENTAJE)) for r in area.itertuples()]
    jrows = [(r.CATEGORIA, fmt_int(r.FILAS), fmt_pct(r.PORCENTAJE)) for r in jornada.itertuples()]
    mrows = [(r.CATEGORIA, fmt_int(r.FILAS), fmt_pct(r.PORCENTAJE)) for r in modalidad.itertuples()]
    add_table(doc, ["Área", "Filas", "%"], arows, [5200, 1900, 2260], font_size=9)
    add_table(doc, ["Jornada", "Filas", "%"], jrows, [5200, 1900, 2260], font_size=9)
    add_table(doc, ["Modalidad", "Filas", "%"], mrows, [5200, 1900, 2260], font_size=9)
    add_body(
        doc,
        "La clasificación urbana representa 79.90% y la rural 20.07%; tres registros no informan el área. Las jornadas más frecuentes son Doble (31.93%), Vespertina (28.82%) y Matutina (25.36%). El 9.24% no informa jornada. La modalidad Monolingüe representa 95.92%.",
    )

    add_heading(doc, "7.3 Ubicación y teléfonos", 2)
    urows = [(r.CATEGORIA, fmt_int(r.FILAS), fmt_pct(r.PORCENTAJE)) for r in ubicacion.itertuples()]
    trows = [(r.CATEGORIA, fmt_int(r.FILAS), fmt_pct(r.PORCENTAJE)) for r in telefono.itertuples()]
    add_table(doc, ["Ubicación", "Filas", "%"], urows, [5200, 1900, 2260], font_size=9)
    add_table(doc, ["Estado del teléfono", "Filas", "%"], trows, [5200, 1900, 2260], font_size=9)
    add_body(
        doc,
        "Ciudad Capital concentra 2,161 registros (18.62%). Después de integrarla a Guatemala, el departamento analítico de Guatemala suma 3,805 filas (32.79%). En teléfonos, 90.01% contiene un único número válido, 1.05% varios números válidos, 0.90% un formato no estándar que debe revisarse y 8.04% está vacío.",
    )

    add_heading(doc, "8. Control de calidad y reproducibilidad", 1)
    def control_value(value):
        try:
            number = float(value)
            return f"{int(number):,}" if number.is_integer() else f"{number:,.4f}"
        except (TypeError, ValueError):
            return str(value)

    control_rows = [(r.CONTROL_ID, r.CONTROL, "Sí" if r.CUMPLE else "No", control_value(r.ESPERADO), control_value(r.OBSERVADO)) for r in controls.itertuples()]
    add_table(doc, ["ID", "Control", "Cumple", "Esperado", "Observado"], control_rows, [720, 4160, 1000, 1680, 1800], font_size=8.1)
    add_body(
        doc,
        "Los 22 controles pasaron. Además, el notebook conserva sus salidas de ejecución: 15 celdas de código tienen número de ejecución consecutivo, no existen trazas de error y las tablas, decisiones y figuras corresponden a la misma corrida que generó los archivos finales.",
    )
    add_heading(doc, "8.1 Consistencia entre variables", 2)
    cross_rows = [(r.VALIDACION_ID, r.VALIDACION, fmt_int(r.INCONSISTENCIAS), "Sí" if r.CUMPLE else "No") for r in cruzadas.itertuples()]
    add_table(doc, ["ID", "Validación cruzada", "Inconsistencias", "Cumple"], cross_rows, [800, 5860, 1500, 1200], font_size=8.4)
    add_code(doc, """
assert controles["CUMPLE"].all()
assert len(df_limpio) == len(df_raw) == 11603
assert df_limpio["CODIGO"].is_unique
assert df_limpio["DEPARTAMENTO_ANALISIS"].nunique() == 22
assert candidatos_duplicados["DECISION"].notna().all()
assert validaciones_cruzadas["INCONSISTENCIAS"].eq(0).all()
""")

    add_heading(doc, "9. Libro de códigos y entregables", 1)
    add_body(
        doc,
        f"El libro de códigos documenta {len(codebook)} variables. Para cada una indica significado, tipo, admisión de null, dominio, valores posibles, tratamiento, origen, método de cálculo, utilidad, fecha de extracción, fuente y versión. Los catálogos de codificación se incluyen en CSV y Excel. Todos los códigos categóricos son nominales y sus etiquetas legibles se conservan.",
    )
    deliverables = [
        ("Proyecto_1_MINEDUC_2026_Final_Ejecutado.ipynb", "Notebook reproducible, ejecutado y con análisis 2026."),
        ("Informe_Final_Proyecto_1_MINEDUC.docx / .pdf", "Informe final formal y versión fija para entrega."),
        ("Libro_de_Codigos_MINEDUC.docx / .pdf / .md", "Descripción general, 59 variables, tratamientos, derivaciones y catálogos."),
        ("Proyecto_1_MINEDUC_Datos_Libro_Codigos.xlsx", "Panel, datos analíticos, diagnóstico, reglas, catálogos y controles."),
        ("establecimientos_diversificado_limpio.csv", "Conjunto nacional completo de 11,603 × 59."),
        ("establecimientos_diversificado_limpio_analitico.csv", "Vista de 11,603 × 42 para uso práctico."),
        ("datos_crudos_establecimientos_diversificado.csv", "Consolidado crudo inalterado."),
        ("bitacora_cambios_detalle.csv", "Trazabilidad por fila y variable."),
        ("catalogos_codificacion.csv", "Relación de valores, etiquetas y códigos nominales."),
        ("candidatos_duplicados_parciales.csv", "13,031 pares, similitud, evidencia, decisión y justificación."),
        ("calidad_antes_despues.csv", "Doce métricas de calidad comparables."),
        ("catalogo_territorial_observado.csv", "334 municipios presentes y 22 zonas capitalinas."),
        ("resumen_transformaciones.csv", "Tabla exigida de cinco columnas para las 17 variables."),
    ]
    add_table(doc, ["Archivo", "Finalidad"], deliverables, [4300, 5060], font_size=8.5)

    add_heading(doc, "10. Limitaciones", 1)
    limitations = [
        "El archivo representa códigos de servicio educativo y no necesariamente planteles físicos únicos.",
        "Las cifras describen el estado de la fuente en la fecha de procesamiento; el portal puede cambiar posteriormente.",
        "No se imputaron nombres, teléfonos, distritos ni jornadas sin evidencia; null conserva la ausencia.",
        "Los 104 teléfonos no estándar quedan como null analítico y requieren revisión manual o contraste con una fuente oficial; el valor original no se pierde.",
        "El catálogo territorial observado cubre 334 de los 340 municipios nacionales reportados por SEGEPLAN; los otros seis no aparecen con Diversificado en este snapshot.",
        "Se recomiendan 2,063 revisiones manuales entre los pares parciales; no se efectuaron fusiones sin evidencia oficial.",
        "La normalización ortográfica mejora presentación, pero un nombre propio crítico debe verificarse contra el registro oficial antes de uso legal.",
        "Las particiones por origen se derivaron del consolidado crudo disponible; su finalidad es auditar la unión nacional.",
    ]
    bullet_id = create_numbering(doc, bullet=True)
    for item in limitations:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "11. Conclusiones", 1)
    add_body(
        doc,
        "El proyecto cumple la ruta completa solicitada en la guía 2026: conserva los datos crudos en CSV, diagnostica su estructura y calidad mediante código, ejecuta reglas justificadas, documenta cada cambio, analiza duplicados exactos y parciales, verifica consistencia entre variables, compara calidad antes/después, produce una unión nacional limpia y entrega un libro de códigos multiformato. La decisión de mantener originales y derivados permite mejorar legibilidad sin perder evidencia.",
    )
    add_body(
        doc,
        "La codificación nominal de variables categóricas —incluida la distinción Ciudad Capital/Otros municipios— reduce inconsistencias futuras. El tratamiento variable por variable de los faltantes evita convertir categorías semánticas de ausencia en valores reales o imputar información inexistente.",
    )

    add_heading(doc, "Referencia", 1)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    r = p.add_run("Ministerio de Educación de Guatemala. (s. f.). Búsqueda de centros educativos autorizados por el Mineduc. ")
    set_run_font(r)
    add_hyperlink(p, "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/", "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    r = p.add_run("Secretaría de Planificación y Programación de la Presidencia. (2026). Cálculo matemático para la asignación constitucional a las municipalidades 2026. ")
    set_run_font(r)
    add_hyperlink(p, "https://datos.segeplan.gob.gt/es/dataset/calculo-matematico-para-la-asignacion-constitucional-a-las-municipalidades-2026", "https://datos.segeplan.gob.gt/es/dataset/calculo-matematico-para-la-asignacion-constitucional-a-las-municipalidades-2026")

    doc.save(REPORT)


def group_for(variable: str) -> str:
    if variable in {"CODIGO", "FUENTE", "FUENTE_URL", "FECHA_EXTRACCION", "FECHA_PROCESAMIENTO", "VERSION_CONJUNTO"}:
        return "Identificación y trazabilidad"
    if variable.startswith("DISTRITO") or variable.startswith("DEPARTAMENTO") or variable.startswith("MUNICIPIO") or variable in {"UBICACION_GRUPO", "TIPO_UBICACION", "ES_CIUDAD_CAPITAL", "ZONA_CAPITAL_COD"}:
        return "Geografía"
    if variable.startswith("ESTABLECIMIENTO") or variable.startswith("DIRECCION"):
        return "Establecimiento y ubicación textual"
    if variable.startswith("TELEFONO"):
        return "Contacto"
    if variable.startswith("SUPERVISOR") or variable.startswith("DIRECTOR"):
        return "Personas responsables"
    return "Características del servicio"


def build_codebook() -> None:
    cb = pd.read_csv(DATA / "libro_codigos.csv").fillna("")
    catalogs = pd.read_csv(DATA / "catalogos_codificacion.csv").fillna("")
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc, "Proyecto 1 · Libro de códigos")
    doc.core_properties.title = "Libro de códigos: establecimientos educativos MINEDUC"
    doc.core_properties.author = "; ".join(TEAM)
    add_cover(
        doc,
        "Libro de códigos",
        "Conjunto limpio de establecimientos educativos · Nivel Diversificado",
        "PROYECTO 1 · DOCUMENTACIÓN DEL CONJUNTO",
    )

    add_heading(doc, "Descripción general", 1)
    add_body(
        doc,
        "El conjunto documentado contiene 11,603 registros de servicios educativos autorizados para el nivel Diversificado y 59 variables. La fuente es el buscador del Ministerio de Educación de Guatemala. La llave es CODIGO. El archivo no debe interpretarse como un censo de edificios, porque un plantel puede asociarse con más de un servicio autorizado.",
    )
    add_metric_cards(doc, [("Filas", "11,603"), ("Variables", "59"), ("Llave", "CODIGO"), ("Municipios", "334"), ("Versión", "2.0.0")])
    add_callout(
        doc,
        "Convenciones",
        "*_ORIGINAL conserva el valor publicado; *_LIMPIO contiene la presentación analítica; *_COD es un código nominal; *_ESTADO describe calidad o ausencia. Null significa ausencia real o semántica documentada, nunca una imputación.",
    )
    add_heading(doc, "Unidad de observación y cobertura", 2)
    metadata = [
        ("Unidad", "Código de servicio educativo autorizado."),
        ("Cobertura", "Guatemala; 23 categorías de origen, 22 departamentos analíticos, 334 municipios observados y 22 zonas capitalinas."),
        ("Nivel", "Diversificado."),
        ("Fecha de extracción", "10 de julio de 2026; fecha del snapshot crudo preservado."),
        ("Fecha de procesamiento", "17 de julio de 2026."),
        ("Versión", "2.0.0."),
        ("Fuente", "MINEDUC, Búsqueda de centros educativos autorizados."),
        ("Formato canónico", "CSV UTF-8; todas las columnas textuales preservan ceros y guiones."),
        ("Trazabilidad", "Columnas originales, bitácora de cambios, catálogos y controles de calidad."),
    ]
    add_table(doc, ["Elemento", "Descripción"], metadata, [2300, 7060], font_size=9.2)

    add_heading(doc, "Diccionario de variables", 1)
    group_order = [
        "Identificación y trazabilidad",
        "Geografía",
        "Establecimiento y ubicación textual",
        "Contacto",
        "Personas responsables",
        "Características del servicio",
    ]
    cb["GRUPO"] = cb["VARIABLE"].map(group_for)
    for group in group_order:
        part = cb.loc[cb["GRUPO"].eq(group)]
        if part.empty:
            continue
        add_heading(doc, group, 2)
        rows = []
        for r in part.itertuples():
            spec = f"Tipo: {r.TIPO}; null: {r.ADMITE_NULL}.\nDominio: {r.DOMINIO_PERMITIDO}\nValores: {r.VALORES_POSIBLES}\nEjemplo: {r.EJEMPLO}"
            derivation = f"Derivada: {r.ES_DERIVADA}.\nOrigen: {r.VARIABLES_ORIGEN}.\nMétodo: {r.METODO_CALCULO}\nUso: {r.UTILIDAD}.\nTratamiento: {r.TRATAMIENTO_LIMPIEZA}"
            rows.append((r.VARIABLE, r.DESCRIPCION, spec, derivation))
        add_table(doc, ["Variable", "Significado", "Tipo, dominio y ejemplo", "Tratamiento y derivación"], rows, [1750, 2700, 2600, 2310], font_size=7.1, header_fill="EAF2F8")

    add_heading(doc, "Catálogos de codificación", 1)
    add_body(
        doc,
        "Los códigos siguientes son nominales, salvo los códigos territoriales cuya estructura proviene de la identificación geográfica. Los números de una categoría no expresan prioridad ni intensidad. Para presentación deben utilizarse las etiquetas limpias.",
    )
    for variable in catalogs["VARIABLE"].drop_duplicates():
        part = catalogs.loc[catalogs["VARIABLE"].eq(variable)]
        add_heading(doc, str(variable), 2)
        rows = [(r.VALOR_ORIGINAL, r.ETIQUETA_LIMPIA, str(r.CODIGO), r.TIPO_CODIGO) for r in part.itertuples()]
        add_table(doc, ["Valor original", "Etiqueta limpia", "Código", "Tipo"], rows, [2500, 2700, 1300, 2860], font_size=8.3)

    add_heading(doc, "Reglas de uso", 1)
    bullet_id = create_numbering(doc, bullet=True)
    rules = [
        "Usar CODIGO como llave del servicio; no usar ESTABLECIMIENTO_LIMPIO para deduplicar.",
        "Conservar las columnas *_ORIGINAL en procesos auditables y reportes sensibles.",
        "Usar DEPARTAMENTO_ANALISIS para comparaciones entre 22 departamentos y DEPARTAMENTO_ORIGINAL para reproducir la fuente.",
        "Interpretar null junto con *_ESTADO o *_INFORMADA cuando exista esa columna.",
        "No tratar códigos nominales como escalas ordinales ni calcular promedios con ellos.",
        "Revisar TELEFONO_ESTADO antes de utilizar números de contacto.",
        "Consultar candidatos_duplicados_parciales.csv antes de fusionar códigos con nombres iguales o similares.",
        "Mantener FUENTE_URL, FECHA_EXTRACCION y VERSION_CONJUNTO al publicar una derivación.",
    ]
    for item in rules:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "Referencia", 1)
    p = add_body(doc, "Ministerio de Educación de Guatemala. Búsqueda de centros educativos autorizados por el Mineduc. ")
    add_hyperlink(p, "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/", "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/")
    doc.save(CODEBOOK)


def build_codebook_markdown() -> None:
    cb = pd.read_csv(DATA / "libro_codigos.csv").fillna("")
    catalogs = pd.read_csv(DATA / "catalogos_codificacion.csv").fillna("")
    metadata = pd.read_csv(DATA / "metadatos_conjunto.csv").fillna("")

    lines = [
        "# Libro de códigos — Establecimientos educativos MINEDUC",
        "",
        "**Proyecto 1 · CC3084 Data Science · Semestre II 2026 · Versión 2.0.0**",
        "",
        "## Descripción general",
        "",
        "El conjunto contiene 11,603 códigos de servicio educativo autorizado del nivel Diversificado y 59 variables. "
        "La llave es `CODIGO`. Los campos `*_ORIGINAL` preservan la fuente, los campos `*_LIMPIO`/`*_LIMPIA` son de presentación analítica y los campos `*_COD` son códigos nominales salvo cuando se indique que son territoriales.",
        "",
        "## Metadatos",
        "",
        "| Campo | Valor |",
        "|---|---|",
    ]
    for r in metadata.itertuples(index=False):
        lines.append(f"| {str(r.CAMPO).replace('|', '/')} | {str(r.VALOR).replace('|', '/')} |")

    lines.extend([
        "", "## Diccionario de variables", "",
        "| Variable | Descripción | Tipo / null | Dominio y valores | Tratamiento / derivación | Fuente, fecha y versión |",
        "|---|---|---|---|---|---|",
    ])
    for r in cb.itertuples(index=False):
        values = str(r.VALORES_POSIBLES).replace("|", "/")
        domain = str(r.DOMINIO_PERMITIDO).replace("|", "/")
        treatment = str(r.TRATAMIENTO_LIMPIEZA).replace("|", "/")
        origin = str(r.VARIABLES_ORIGEN).replace("|", "/")
        lines.append(
            f"| `{r.VARIABLE}` | {str(r.DESCRIPCION).replace('|', '/')} | {r.TIPO}; null: {r.ADMITE_NULL} | "
            f"{domain}. Valores: {values} | {treatment} Derivada: {r.ES_DERIVADA}; origen: {origin}; método: {str(r.METODO_CALCULO).replace('|', '/')}; uso: {str(r.UTILIDAD).replace('|', '/')}. | "
            f"{r.FUENTE}; {r.FECHA_EXTRACCION}; v{r.VERSION_CONJUNTO} |"
        )

    lines.extend(["", "## Catálogos de codificación", ""])
    for variable in catalogs["VARIABLE"].drop_duplicates():
        lines.extend([
            f"### {variable}", "",
            "| Valor original | Etiqueta limpia | Código | Tipo |", "|---|---|---:|---|",
        ])
        part = catalogs.loc[catalogs["VARIABLE"].eq(variable)]
        for r in part.itertuples(index=False):
            lines.append(f"| {r.VALOR_ORIGINAL} | {r.ETIQUETA_LIMPIA} | {r.CODIGO} | {r.TIPO_CODIGO} |")
        lines.append("")

    lines.extend([
        "## Reglas de uso", "",
        "1. Usar `CODIGO` como llave; no deduplicar por nombre.",
        "2. Conservar los campos `*_ORIGINAL` en cualquier proceso auditable.",
        "3. Interpretar los códigos categóricos como nominales, no ordinales.",
        "4. Consultar los indicadores de ausencia y el archivo de candidatos parciales antes de tomar decisiones.",
        "5. Mantener fuente, fecha de extracción y versión al publicar derivados.",
        "",
    ])
    CODEBOOK_MD.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    build_report()
    build_codebook()
    build_codebook_markdown()
    print(REPORT)
    print(CODEBOOK)
    print(CODEBOOK_MD)
