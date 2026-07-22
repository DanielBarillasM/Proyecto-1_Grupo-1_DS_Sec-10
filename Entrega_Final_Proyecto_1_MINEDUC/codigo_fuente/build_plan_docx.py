from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Plan_Formal_Proyecto_1_Obtencion_y_Limpieza_MINEDUC.docx"

sys.path.insert(0, str(ROOT))
from table_geometry import apply_table_geometry  # noqa: E402


# Preset: standard_business_brief, with a restrained academic cover override.
NAVY = "12355B"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17202A"
MUTED = "5F6B76"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C4CE"
GOLD = "B98A2E"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, *, color=BORDER, size=6, sides=("top", "left", "bottom", "right")) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        if side in sides or side.startswith("inside"):
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(size))
            element.set(qn("w:color"), color)
        else:
            element.set(qn("w:val"), "nil")


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    """Evita que Word/LibreOffice dividan una fila entre dos páginas."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run_font(run, *, name="Calibri", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_widow_control(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    widow = p_pr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)
    widow.set(qn("w:val"), "1")


def add_field(paragraph, instruction: str, fallback="1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend([r_pr, text_el])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(doc: Document, *, bullet=False) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    set_widow_control(paragraph)


def add_numbered(doc, text: str, num_id: int, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text: str, bullet_id: int, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_id)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_body(doc, text: str, *, bold_lead: str | None = None, italic=False):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    set_widow_control(p)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p, True)
    return p


def add_callout(doc, label: str, text: str, *, fill=LIGHT_BLUE, accent=NAVY):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color=accent, size=14, sides=("left",))
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=accent, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=160,
        cell_margins_dxa={"top": 120, "bottom": 120, "start": 160, "end": 160},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths, *, font_size=9.3, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, header, bold=True, color=NAVY, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_borders(cell)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=font_size, align=align)
            set_cell_borders(cells[idx])
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return table


def add_labeled_paragraph(doc, label: str, value: str, *, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True, color=NAVY)
    r2 = p.add_run(value)
    set_run_font(r2, color=INK)
    return p


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Table Source" not in [s.name for s in doc.styles]:
        source = doc.styles.add_style("Table Source", WD_STYLE_TYPE.PARAGRAPH)
        source.font.name = "Calibri"
        source.font.size = Pt(9)
        source.font.italic = True
        source.font.color.rgb = RGBColor.from_string(MUTED)
        source.paragraph_format.space_before = Pt(4)
        source.paragraph_format.space_after = Pt(4)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("PROYECTO 1  |  PLAN DE OBTENCIÓN Y LIMPIEZA")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("DATA SCIENCE")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("Página ")
    set_run_font(label, size=9, color=MUTED)
    add_field(fp, "PAGE", "1")
    sep = fp.add_run(" de ")
    set_run_font(sep, size=9, color=MUTED)
    add_field(fp, "NUMPAGES", "1")


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "Plan formal del Proyecto 1: Obtención y Limpieza de datos"
    doc.core_properties.subject = "Plan de trabajo para datos de establecimientos educativos del MINEDUC"
    doc.core_properties.author = "Pablo Daniel Barillas Moreno"
    doc.core_properties.keywords = "Data Science, MINEDUC, limpieza de datos, plan de trabajo"

    number_id = create_numbering(doc, bullet=False)
    bullet_id = create_numbering(doc, bullet=True)

    # Cover: editorial_cover pattern, academic override.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(62)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    kr = kicker.add_run("UNIVERSIDAD DEL VALLE DE GUATEMALA")
    set_run_font(kr, size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    tr = title.add_run("Plan de obtención, diagnóstico\ny limpieza de datos")
    set_run_font(tr, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    sr = subtitle.add_run("Establecimientos educativos de Guatemala\nNivel Diversificado")
    set_run_font(sr, size=15, color=BLUE, bold=False)

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(42)
    pr = project.add_run("AVANCE DEL PROYECTO 1  |  DATA SCIENCE")
    set_run_font(pr, size=10.5, color=MUTED, bold=True)

    add_labeled_paragraph(doc, "Estudiante: ", "Pablo Daniel Barillas Moreno", center=True)
    add_labeled_paragraph(doc, "Carné: ", "22193", center=True)
    add_labeled_paragraph(doc, "Facultad: ", "Ingeniería", center=True)
    add_labeled_paragraph(doc, "Fecha: ", "10 de julio de 2026", center=True)

    cover_note = doc.add_paragraph()
    cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_note.paragraph_format.space_before = Pt(36)
    cover_note.paragraph_format.space_after = Pt(0)
    cn = cover_note.add_run("Documento de planificación; no constituye la ejecución final de la limpieza.")
    set_run_font(cn, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()

    add_heading(doc, "Resumen ejecutivo", 1)
    add_body(
        doc,
        "Este documento establece el plan que se seguirá para obtener, diagnosticar, limpiar, unir y documentar los datos de establecimientos educativos de Guatemala que cuentan con servicios autorizados en el nivel Diversificado. El propósito del avance es definir una ruta transparente y reproducible; por tanto, las operaciones de limpieza se presentan como acciones propuestas y no como resultados finales ya ejecutados.",
    )
    add_body(
        doc,
        "El diagnóstico preliminar se apoya en la extracción realizada desde el buscador oficial del Ministerio de Educación. Esta revisión permite dimensionar el trabajo, identificar las variables prioritarias y diseñar validaciones antes de modificar los datos.",
    )
    add_callout(
        doc,
        "Alcance",
        "El avance describe el estado inicial de los datos y la estrategia de trabajo. La limpieza definitiva, la unión final, el archivo limpio y el libro de códigos se producirán en las fases posteriores.",
    )

    metrics = [
        ("Filas crudas identificadas", "11,603"),
        ("Variables", "17"),
        ("Códigos únicos", "11,603"),
        ("Filas exactamente duplicadas", "0"),
        ("Nivel educativo", "Diversificado"),
    ]
    add_table(doc, ["Indicador preliminar", "Resultado"], metrics, [3600, 5760], font_size=10)

    add_heading(doc, "1. Contexto y alcance del proyecto", 1)
    add_body(
        doc,
        "El proyecto parte de información pública disponible en el buscador de centros educativos autorizados del MINEDUC. Cada registro corresponde a un código de servicio educativo y no necesariamente a un edificio físico único. Un establecimiento puede aparecer con más de un código debido a diferencias de jornada, plan u otros servicios autorizados. Esta característica condicionará especialmente la detección de duplicados.",
    )
    add_body(
        doc,
        "El plan abarcará la descarga de los datos, su conservación en formato CSV, el diagnóstico de calidad, la ejecución documentada de reglas de limpieza, la unión nacional y la elaboración de un libro de códigos. En todas las etapas se conservará una copia inalterada de los datos originales.",
    )

    add_heading(doc, "2. Objetivos", 1)
    add_heading(doc, "2.1 Objetivo general", 2)
    add_body(
        doc,
        "Definir un proceso reproducible para obtener, diagnosticar, limpiar, integrar y documentar los datos de establecimientos educativos de Guatemala con nivel Diversificado, preservando la calidad ortográfica de los nombres y la trazabilidad de cada transformación.",
    )
    add_heading(doc, "2.2 Objetivos específicos", 2)
    objectives = [
        "Descargar y conservar los datos crudos correspondientes a todas las categorías geográficas disponibles en el portal del MINEDUC.",
        "Describir la estructura, el volumen y los principales problemas de calidad del conjunto inicial.",
        "Diseñar reglas de limpieza específicas para códigos, nombres, direcciones, teléfonos, personas y variables categóricas.",
        "Evitar la eliminación incorrecta de registros que representen servicios educativos distintos.",
        "Producir un conjunto nacional unificado, validado y listo para análisis.",
        "Elaborar un libro de códigos que documente las variables, los valores posibles y las decisiones de limpieza.",
    ]
    for item in objectives:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "3. Plan de obtención y resguardo de los datos", 1)
    add_heading(doc, "3.1 Fuente oficial", 2)
    p = add_body(doc, "Los datos se consultarán en el buscador oficial del Ministerio de Educación de Guatemala: ")
    add_hyperlink(p, "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/", "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/")
    add_body(
        doc,
        "Se seleccionará el nivel Diversificado y se recorrerán todas las opciones geográficas que ofrece el portal. La fecha y hora de extracción, los filtros utilizados y el número de registros reportado por cada consulta quedarán registrados en un manifiesto de obtención.",
    )

    add_heading(doc, "3.2 Procedimiento propuesto de descarga", 2)
    acquisition_steps = [
        "Acceder al portal y confirmar que la fuente se encuentre disponible.",
        "Seleccionar el nivel escolar Diversificado, sin restringir sector, plan o modalidad.",
        "Consultar cada departamento o categoría geográfica por separado y registrar el total mostrado.",
        "Cuando una consulta supere el límite de visualización del portal, subdividirla por municipio o zona y comprobar que la suma coincida con el total informado.",
        "Exportar o transformar cada resultado a CSV con codificación UTF-8, sin modificar los valores de origen.",
        "Verificar que todos los archivos tengan las mismas columnas y conservar una bitácora de archivos, fechas y conteos.",
    ]
    for item in acquisition_steps:
        add_numbered(doc, item, number_id)

    add_callout(
        doc,
        "Control especial de Ciudad Capital",
        "La consulta preliminar reportó 2,161 registros, pero la vista general mostró solamente 2,000. Por ello se prevé consultar sus 22 zonas, unir los códigos y validar que el total recuperado sea 2,161.",
        fill="FFF7E6",
        accent=GOLD,
    )

    add_heading(doc, "3.3 Organización de los archivos", 2)
    file_structure = [
        ("datos/crudos/", "CSV descargados sin modificación, separados por departamento o filtro del portal."),
        ("datos/intermedios/", "Resultados temporales de normalización y listas de revisión manual."),
        ("datos/limpios/", "Conjunto nacional validado y listo para análisis."),
        ("documentacion/", "Bitácora de cambios, manifiesto de obtención y libro de códigos."),
        ("notebooks/", "Cuadernos de carga, diagnóstico, limpieza, unión y validación."),
    ]
    add_table(doc, ["Ubicación", "Contenido previsto"], file_structure, [2400, 6960], font_size=9.8)

    add_heading(doc, "4. Descripción preliminar del conjunto crudo", 1)
    add_body(
        doc,
        "La extracción preliminar contiene 11,603 filas y 17 variables. Los 11,603 códigos son únicos y no se detectaron filas completamente idénticas. El conjunto incluye registros abiertos, cerrados temporalmente, cerrados definitivamente y estados temporales; estos estados se conservarán porque constituyen información analítica y no errores de calidad.",
    )
    add_body(
        doc,
        "El portal utiliza 23 categorías geográficas porque separa Ciudad Capital de Guatemala. En la integración final se conservará la clasificación original y, si se necesita análisis territorial por departamento, se creará una variable derivada que ubique Ciudad Capital dentro del departamento de Guatemala.",
    )

    variable_rows = [
        ("CODIGO", "Identificador del servicio educativo; integra ubicación, correlativo y nivel."),
        ("DISTRITO", "Código del distrito o supervisión educativa asociada."),
        ("DEPARTAMENTO", "Categoría geográfica del portal; Ciudad Capital aparece separada."),
        ("MUNICIPIO", "Municipio o zona de Ciudad Capital."),
        ("ESTABLECIMIENTO", "Nombre registrado del establecimiento educativo."),
        ("DIRECCION", "Dirección textual del establecimiento."),
        ("TELEFONO", "Uno o varios números telefónicos reportados."),
        ("SUPERVISOR", "Nombre de la persona supervisora responsable."),
        ("DIRECTOR", "Nombre de la persona directora reportada."),
        ("NIVEL", "Nivel escolar consultado; será Diversificado."),
        ("SECTOR", "Sector oficial, privado, municipal o por cooperativa."),
        ("AREA", "Clasificación urbana, rural o sin especificar."),
        ("STATUS", "Estado administrativo del código educativo."),
        ("MODALIDAD", "Modalidad monolingüe o bilingüe."),
        ("JORNADA", "Jornada autorizada del servicio."),
        ("PLAN", "Plan temporal o modalidad de atención."),
        ("DEPARTAMENTAL", "Dirección departamental de educación responsable."),
    ]
    add_table(doc, ["Variable", "Descripción preliminar"], variable_rows, [2300, 7060], font_size=9.2)
    source = doc.add_paragraph("Fuente: elaboración propia a partir de los campos publicados por el MINEDUC.", style="Table Source")

    add_heading(doc, "5. Estado preliminar de los datos", 1)
    add_body(
        doc,
        "El diagnóstico inicial se utilizará para priorizar las operaciones. Los conteos siguientes no representan una limpieza concluida; únicamente cuantifican vacíos, marcadores y variantes candidatas que deberán revisarse.",
    )
    priority_rows = [
        ("ESTABLECIMIENTO", "5 vacíos; 947 grupos candidatos de variantes que afectan 4,925 filas.", "Ortografía, tildes, siglas, puntuación y repetición de nombres."),
        ("DIRECTOR", "2,106 faltantes efectivos; 146 grupos candidatos.", "Vacíos, guiones, SIN DATO y diferencias de acentuación."),
        ("DIRECCION", "87 faltantes efectivos; 576 grupos candidatos.", "Abreviaturas, puntuación y expresiones distintas de una ubicación."),
        ("TELEFONO", "933 vacíos; 250 valores no estándar.", "Múltiples números, separadores y longitudes distintas de ocho dígitos."),
        ("SUPERVISOR", "525 vacíos; 183 grupos candidatos.", "Variantes ortográficas y de acentuación en nombres de personas."),
        ("DISTRITO", "591 valores vacíos o incompletos.", "Campos sin valor o formas incompletas como 01-."),
        ("JORNADA", "1,072 registros con SIN JORNADA.", "Categoría informativa que debe distinguirse de un nulo técnico."),
    ]
    add_table(doc, ["Variable prioritaria", "Evidencia preliminar", "Aspecto que se revisará"], priority_rows, [2050, 3160, 4150], font_size=8.8)

    add_heading(doc, "5.1 Aspectos que no se tratarán como errores", 2)
    non_errors = [
        "Los distintos valores de STATUS, incluidos los cierres, se conservarán.",
        "Los códigos diferentes asociados con nombres iguales no se eliminarán automáticamente.",
        "Los planes, jornadas y modalidades autorizadas se mantendrán como categorías válidas.",
        "SIN JORNADA y SIN ESPECIFICAR se distinguirán de los campos vacíos.",
    ]
    for item in non_errors:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "6. Estrategia general de limpieza", 1)
    principles = [
        ("Preservar el origen. ", "Los CSV crudos no se sobrescribirán; la limpieza se realizará sobre copias o nuevas columnas."),
        ("Mantener la ortografía. ", "No se eliminarán tildes, diéresis ni la letra ñ de los nombres destinados a informes."),
        ("Separar comparación y presentación. ", "Podrá utilizarse una clave auxiliar en mayúsculas y sin signos para detectar candidatos, pero nunca como texto final."),
        ("Documentar cada cambio. ", "Toda sustitución tendrá un registro con valor original, valor propuesto, regla, razón y filas afectadas."),
        ("No inventar datos. ", "Los teléfonos, distritos, nombres o direcciones faltantes no se completarán sin una fuente oficial."),
        ("Validar después de cada etapa. ", "Se compararán filas, códigos únicos, faltantes, categorías y formatos antes y después."),
    ]
    for lead, rest in principles:
        add_bullet(doc, lead + rest, bullet_id, bold_lead=lead)

    add_callout(
        doc,
        "Principio ortográfico",
        "La normalización informática servirá para encontrar diferencias, no para borrar reglas del español. La versión limpia deberá conservar una escritura apta para informes institucionales.",
    )

    add_heading(doc, "7. Estrategia específica por variable", 1)
    add_heading(doc, "7.1 Nombre del establecimiento", 2)
    establishment_steps = [
        "Conservar ESTABLECIMIENTO_ORIGINAL sin cambios y crear ESTABLECIMIENTO_LIMPIO.",
        "Aplicar normalización Unicode NFC para uniformar la representación informática de caracteres.",
        "Eliminar únicamente espacios al inicio, al final y repeticiones internas innecesarias.",
        "Estandarizar de manera controlada comillas, puntos y guiones alrededor de siglas.",
        "Detectar variantes que cambien solo por tildes o puntuación, como EDUCACION y EDUCACIÓN.",
        "Mantener siglas institucionales —por ejemplo, INED, INEB, CEEX e IGER— en mayúsculas.",
        "Evitar el uso automático de formato tipo título, ya que podría alterar siglas, apellidos o números romanos.",
        "Construir un diccionario de correcciones verificadas y aplicarlo de forma reproducible.",
        "Utilizar similitud difusa únicamente para generar candidatos de revisión, nunca para reemplazar directamente.",
        "Revisar manualmente los casos ambiguos y registrar la decisión tomada.",
        "No eliminar una fila por repetir el nombre; antes se compararán código, ubicación, dirección, jornada y plan.",
    ]
    establishment_number_id = create_numbering(doc, bullet=False)
    for item in establishment_steps:
        add_numbered(doc, item, establishment_number_id)

    add_heading(doc, "7.2 Director y supervisor", 2)
    for item in [
        "Convertir vacíos, secuencias de guiones y expresiones como SIN DATO a una representación uniforme de faltante en la copia limpia.",
        "Conservar el valor original para auditoría.",
        "Corregir tildes únicamente mediante reglas verificadas o revisión manual.",
        "No fusionar personas solamente porque sus nombres sean similares; pueden existir homónimos.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "7.3 Dirección", 2)
    for item in [
        "Normalizar Unicode, espacios y puntuación básica sin eliminar referencias geográficas.",
        "Definir reglas consistentes para abreviaturas como AV., AVENIDA, CALLE y KM.",
        "Conservar números, zonas, aldeas, colonias, lotes y puntos de referencia.",
        "Marcar como faltantes las direcciones vacías o compuestas únicamente por signos.",
        "Comparar direcciones similares dentro del mismo municipio como apoyo a la revisión, no como criterio automático de eliminación.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "7.4 Teléfono", 2)
    for item in [
        "Mantener el campo como texto para conservar ceros iniciales.",
        "Identificar celdas con varios teléfonos separados por guiones, comas, espacios o barras.",
        "Extraer y validar cada candidato contra el formato guatemalteco actual de ocho dígitos.",
        "No completar números antiguos, cortos o incompletos sin una fuente de verificación.",
        "Conservar TELEFONO_ORIGINAL y producir una representación limpia; si hay varios números, utilizar una lista normalizada o una tabla secundaria.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "7.5 Código y distrito", 2)
    for item in [
        "Usar CODIGO como llave del servicio educativo, debido a que los valores preliminares son únicos y cumplen el patrón esperado.",
        "Mantener CODIGO y DISTRITO como texto para proteger ceros iniciales.",
        "Validar el formato de DISTRITO y representar vacíos o formas incompletas como faltantes estandarizados.",
        "No deducir DISTRITO desde CODIGO sin contrastarlo con una fuente oficial.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "7.6 Variables geográficas y categóricas", 2)
    for item in [
        "Crear catálogos permitidos para departamento, municipio, sector, área, estado, modalidad, jornada y plan.",
        "Aplicar mapeos ortográficos explícitos, por ejemplo PETEN → Petén, QUICHE → Quiché y SOLOLA → Sololá.",
        "Conservar CIUDAD CAPITAL como valor original del portal y crear, si se requiere, una variable territorial derivada con departamento Guatemala.",
        "No convertir SIN JORNADA o SIN ESPECIFICAR en vacíos sin conservar su significado original.",
        "Registrar cualquier categoría inesperada para revisión antes de aceptarla o modificarla.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "8. Tratamiento de faltantes y duplicados", 1)
    add_heading(doc, "8.1 Valores faltantes", 2)
    add_body(
        doc,
        "Se elaborará una lista explícita de marcadores equivalentes a ausencia de información, por ejemplo cadenas vacías, guiones y SIN DATO. Estos valores se convertirán a una representación uniforme en la copia limpia, pero su forma original permanecerá disponible en la bitácora o en las columnas originales. La imputación solo se realizará cuando exista una regla confiable y una fuente verificable.",
    )
    add_heading(doc, "8.2 Duplicados", 2)
    duplicate_steps = [
        "Detectar filas completamente idénticas y reportar su cantidad.",
        "Verificar la unicidad de CODIGO antes y después de la limpieza.",
        "Identificar nombres repetidos como candidatos, no como duplicados confirmados.",
        "Comparar departamento, municipio, dirección, jornada, plan y modalidad para comprender por qué se repite un establecimiento.",
        "Eliminar únicamente registros cuya duplicidad pueda demostrarse y documentar la razón.",
        "Generar un reporte de casos conservados, fusionados o eliminados.",
    ]
    duplicate_number_id = create_numbering(doc, bullet=False)
    for item in duplicate_steps:
        add_numbered(doc, item, duplicate_number_id)

    add_heading(doc, "9. Flujo reproducible de trabajo", 1)
    workflow = [
        ("1. Ingesta", "Cargar todos los CSV como texto y conservar códigos y teléfonos sin pérdida de ceros.", "Tabla de archivos y conteos."),
        ("2. Diagnóstico", "Calcular estructura, vacíos, marcadores, valores únicos, formatos y candidatos de duplicidad.", "Perfil de calidad reproducible."),
        ("3. Normalización", "Aplicar funciones por variable sobre una copia y registrar cada regla.", "Datos intermedios + bitácora."),
        ("4. Revisión", "Evaluar variantes ortográficas y casos ambiguos sin reemplazos automáticos.", "Diccionario validado."),
        ("5. Integración", "Verificar esquemas y concatenar los departamentos y zonas.", "Conjunto nacional unificado."),
        ("6. Control final", "Ejecutar pruebas de integridad, formato, categorías y conservación de filas.", "CSV limpio aprobado."),
        ("7. Documentación", "Completar el libro de códigos y resumir decisiones y limitaciones.", "PDF del libro de códigos."),
    ]
    add_table(doc, ["Fase", "Actividad principal", "Evidencia o salida"], workflow, [1750, 5200, 2410], font_size=8.9)

    add_heading(doc, "10. Unión de los datos de todos los departamentos", 1)
    integration_steps = [
        "Comprobar que todos los archivos tengan las mismas 17 columnas y el mismo orden.",
        "Estandarizar nombres de columnas sin modificar los valores crudos.",
        "Concatenar los archivos y registrar el número de filas aportado por cada fuente.",
        "Integrar Ciudad Capital a partir de sus zonas cuando la consulta general esté truncada.",
        "Verificar que la suma de filas por archivo coincida con el total nacional.",
        "Confirmar que CODIGO permanezca único y que no aparezcan duplicados exactos inesperados.",
        "Guardar el archivo unificado intermedio y, después de aplicar las reglas, el conjunto limpio final.",
    ]
    integration_number_id = create_numbering(doc, bullet=False)
    for item in integration_steps:
        add_numbered(doc, item, integration_number_id)

    add_heading(doc, "11. Plan para el libro de códigos", 1)
    add_body(
        doc,
        "El libro de códigos se elaborará después de cerrar las reglas de limpieza. Incluirá una descripción general del conjunto, su unidad de observación, fuente, fecha de extracción, alcance geográfico, número final de filas y variables, y limitaciones conocidas.",
    )
    codebook_fields = [
        ("Nombre de la variable", "Nombre definitivo utilizado en el CSV limpio."),
        ("Definición", "Significado de la variable y su relación con el registro educativo."),
        ("Tipo de dato", "Texto, categoría, identificador u otro tipo aplicable."),
        ("Valores posibles", "Catálogo o patrón permitido."),
        ("Faltantes", "Forma en que se representará la ausencia de información."),
        ("Transformaciones", "Reglas aplicadas desde el valor crudo."),
        ("Ejemplo", "Valor ilustrativo válido."),
        ("Observaciones", "Excepciones, restricciones y notas de interpretación."),
    ]
    add_table(doc, ["Campo del libro", "Contenido esperado"], codebook_fields, [2500, 6860], font_size=9.5)

    add_heading(doc, "12. Controles de calidad y criterios de aceptación", 1)
    acceptance = [
        "El número de filas crudas deberá coincidir con la suma documentada de los archivos de origen.",
        "CODIGO deberá conservarse como texto, cumplir su patrón y permanecer único.",
        "Las 17 variables esperadas deberán estar presentes o sus cambios deberán documentarse en el libro de códigos.",
        "Ninguna regla podrá eliminar tildes, ñ o signos necesarios de la versión destinada a informes.",
        "Toda eliminación de una fila requerirá evidencia de duplicidad y registro en la bitácora.",
        "Las categorías deberán pertenecer a catálogos permitidos o quedar señaladas para revisión.",
        "Los valores faltantes deberán estar representados de forma consistente sin imputaciones no justificadas.",
        "El CSV final deberá poder cargarse nuevamente sin errores de codificación, delimitación o tipos.",
        "El notebook o script deberá ejecutarse desde los datos crudos hasta la salida final sin pasos manuales ocultos.",
    ]
    for item in acceptance:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "13. Entregables previstos", 1)
    deliverables = [
        ("CSV crudos", "Archivos separados por fuente geográfica, sin modificaciones."),
        ("Manifiesto de obtención", "Fuente, fecha, filtros, conteos y observaciones por archivo."),
        ("Notebook o script", "Código reproducible de carga, diagnóstico, limpieza, unión y validación."),
        ("Bitácora de limpieza", "Regla, justificación, variable y cantidad de filas afectadas."),
        ("Diccionario de correcciones", "Mapeos ortográficos y categóricos verificados."),
        ("CSV limpio nacional", "Unión validada de todos los departamentos y zonas."),
        ("Libro de códigos en PDF", "Descripción general, variables, tipos, valores posibles y transformaciones."),
    ]
    add_table(doc, ["Entregable", "Contenido"], deliverables, [2600, 6760], font_size=9.5)

    add_heading(doc, "14. Cronograma propuesto", 1)
    schedule = [
        ("Fase 1", "Obtención, resguardo y manifiesto de los CSV crudos.", "Datos crudos controlados."),
        ("Fase 2", "Perfil de calidad y priorización de variables.", "Diagnóstico del avance."),
        ("Fase 3", "Implementación de funciones de limpieza y bitácora.", "Datos intermedios."),
        ("Fase 4", "Revisión ortográfica y resolución de casos ambiguos.", "Diccionario aprobado."),
        ("Fase 5", "Unión nacional y controles de integridad.", "CSV limpio."),
        ("Fase 6", "Libro de códigos, revisión final y empaquetado.", "Entrega completa."),
    ]
    add_table(doc, ["Fase", "Trabajo principal", "Resultado"], schedule, [1700, 5260, 2400], font_size=9.2)
    add_body(
        doc,
        "Para la entrega del avance del 16 de julio de 2026 se presentarán el diagnóstico preliminar y este plan. Las fases de limpieza definitiva, unión final y libro de códigos se completarán conforme al calendario general del proyecto.",
    )

    add_heading(doc, "15. Riesgos y medidas de mitigación", 1)
    risks = [
        ("Límite de filas del portal", "Comparar total reportado con filas obtenidas y subdividir consultas por municipio o zona."),
        ("Cambios en la fuente", "Registrar fecha, filtros y estructura; conservar copias crudas inmutables."),
        ("Correcciones ortográficas erróneas", "Aplicar diccionarios verificados y revisión manual de casos ambiguos."),
        ("Eliminación de servicios legítimos", "Usar CODIGO como llave y no deduplicar únicamente por nombre."),
        ("Pérdida de ceros iniciales", "Cargar códigos y teléfonos explícitamente como texto."),
        ("Imputación sin evidencia", "Conservar faltantes y completar solo con una fuente oficial documentada."),
    ]
    add_table(doc, ["Riesgo", "Medida de mitigación"], risks, [2800, 6560], font_size=9.4)

    add_heading(doc, "16. Conclusión", 1)
    add_body(
        doc,
        "El plan propuesto permitirá desarrollar el proyecto de forma gradual, verificable y reproducible. El diagnóstico preliminar confirma que la mayor carga de limpieza se concentrará en los nombres de establecimientos y personas, las direcciones, los teléfonos y los distritos. Sin embargo, la prioridad no será reducir filas, sino mejorar la consistencia sin destruir información legítima.",
    )
    add_body(
        doc,
        "La separación entre datos originales, claves auxiliares de comparación y valores limpios de presentación será esencial para cumplir la exigencia de conservar la ortografía. De esta manera, el conjunto final podrá utilizarse tanto para análisis de datos como para la elaboración posterior de informes institucionales.",
    )

    add_heading(doc, "Referencia", 1)
    ref = doc.add_paragraph(style="Normal")
    ref.paragraph_format.left_indent = Inches(0.35)
    ref.paragraph_format.first_line_indent = Inches(-0.35)
    ref.paragraph_format.space_after = Pt(6)
    rr = ref.add_run("Ministerio de Educación de Guatemala. (s. f.). Búsqueda de centros educativos autorizados por el Mineduc. ")
    set_run_font(rr)
    add_hyperlink(ref, "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/", "http://www.mineduc.gob.gt/BUSCAESTABLECIMIENTO_GE/")

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
